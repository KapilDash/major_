"""API Routes - Chatbot with Document Upload Support"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from app.models.schemas import ChatRequest, ChatResponse
from app.services.google_api_service import GoogleAPIClient
import logging
import uuid
import re

logger = logging.getLogger(__name__)
router = APIRouter()
google_client = GoogleAPIClient()

# In-memory document store keyed by conversation_id. Each conversation can contain
# multiple documents so a case can be reviewed as one record.
document_store = {}


def extract_case_metadata(text: str) -> dict:
    """Extract labelled case metadata without treating headings as party names."""
    lines = [re.sub(r"\s+", " ", line).strip(" :-,;") for line in text.splitlines()]
    lines = [line for line in lines if line]

    def first_value(patterns, default):
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                value = re.sub(r"\s+", " ", match.group(1)).strip(" :-,;.")
                if value and value.lower() not in {"case file", "case details", "not available", "n/a"}:
                    return value[:160]
        return default

    plaintiff = first_value([r"^\s*plaintiff\s*(?:/|and)?\s*(?:petitioner)?\s*[:-]\s*([^\r\n]+)", r"^\s*petitioner\s*[:-]\s*([^\r\n]+)", r"^\s*appellant\s*[:-]\s*([^\r\n]+)"], "Not detected")
    defendant = first_value([r"^\s*defendant\s*(?:/|and)?\s*(?:respondent)?\s*[:-]\s*([^\r\n]+)", r"^\s*respondent\s*[:-]\s*([^\r\n]+)", r"^\s*accused\s*[:-]\s*([^\r\n]+)"], "Not detected")
    versus_index = next((index for index, line in enumerate(lines) if re.fullmatch(r"versus|vs\.?", line, re.IGNORECASE)), None)
    if versus_index is not None:
        before = [line for line in lines[:versus_index] if not re.search(r"complaint|petition|case file|commission|court|no\.\s*\d", line, re.IGNORECASE)]
        after = [line for line in lines[versus_index + 1:] if not re.search(r"complaint|petition|case file|commission|court|no\.\s*\d", line, re.IGNORECASE)]
        if before:
            plaintiff = before[-1]
        if after:
            defendant = after[0]
    court = first_value([r"^\s*court\s*(?:name)?\s*[:-]\s*([^\r\n]+)", r"^\s*case filed in\s*[:-]\s*([^\r\n]+)"], "Not detected")
    if court == "Not detected":
        court = next((line for line in lines if re.search(r"commission|tribunal|high court|supreme court|district court", line, re.IGNORECASE)), court)
    return {
        "plaintiff": plaintiff,
        "defendant": defendant,
        "judge": first_value([r"^\s*(?:hon'?ble\s+)?judge\s*[:-]\s*([^\r\n]+)", r"^\s*(?:presiding|president|member|before)\s*(?:officer|member)?\s*[:-]\s*([^\r\n]+)"], "Not detected"),
        "court": court,
        "state": first_value([r"^\s*state\s*[:-]\s*([^\r\n]+)"], "Not detected"),
    }


def extract_pdf_text(content: bytes) -> str:
    """Extract embedded PDF text using the most reliable available parser."""
    import io

    extracted_text = ""
    try:
        import fitz

        document = fitz.open(stream=content, filetype="pdf")
        extracted_text = "\n\n".join(page.get_text("text") for page in document)
        document.close()
    except ImportError:
        pass
    except Exception as error:
        logger.warning("PyMuPDF could not read PDF: %s", error)

    if extracted_text.strip():
        return extracted_text

    try:
        import PyPDF2

        reader = PyPDF2.PdfReader(io.BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        raise HTTPException(status_code=503, detail="PDF extraction is unavailable. Install backend requirements and restart the server.")
    except Exception as error:
        logger.warning("PyPDF2 could not read PDF: %s", error)
        return ""


@router.post("/ask", response_model=dict)
async def ask_chatbot(request: ChatRequest):
    """Ask legal question to chatbot, optionally with document context"""
    try:
        # Check if there's uploaded document context for this conversation
        doc_context = None
        if request.conversation_id and request.conversation_id in document_store:
            documents = document_store[request.conversation_id]
            doc_context = {
                "filename": ", ".join(document["filename"] for document in documents),
                "size": sum(document["size"] for document in documents),
                "text": "\n\n--- NEXT CASE DOCUMENT ---\n\n".join(document["text"] for document in documents),
                "char_count": sum(document["char_count"] for document in documents),
                "document_count": len(documents),
            }

        # The frontend persists case text locally, so chat still works after a backend restart.
        if not doc_context and request.context and request.context.get("case_data"):
            case_text = str(request.context["case_data"])
            doc_context = {
                "filename": "saved case record",
                "size": len(case_text),
                "text": case_text[:15000],
                "char_count": len(case_text[:15000]),
                "document_count": 0,
            }

        # Generate response using Google API
        response = await google_client.generate_chatbot_response(
            query=request.message,
            context=request.context.get("case_data") if request.context else None,
            document_context=doc_context,
            chat_history=request.context.get("chat_history", []) if request.context else [],
        )
        
        return {
            "message_id": str(uuid.uuid4()),
            "content": response.get("response", ""),
            "verified": response.get("verified", False),
            "references": response.get("references", []),
            "confidence": 0.85,
            "source": response.get("source", "Legal Knowledge Base"),
            "document_context_active": doc_context is not None,
            "quota_notice": response.get("quota_notice"),
        }
    except Exception as e:
        logger.error(f"Error in ask_chatbot: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    conversation_id: str = Form(default=None)
):
    """Upload a document for case-specific Q&A"""
    try:
        if not conversation_id:
            conversation_id = str(uuid.uuid4())
        
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        filename = file.filename.lower()
        text = ""
        
        if filename.endswith(".txt"):
            text = content.decode("utf-8", errors="ignore")
        elif filename.endswith(".pdf"):
            text = extract_pdf_text(content)
        elif filename.endswith((".doc", ".docx")):
            # Basic text extraction
                if filename.endswith('.docx'):
                    try:
                        import io
                        from docx import Document
                        document = Document(io.BytesIO(content))
                        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
                        for table in document.tables:
                            text += "\n" + "\n".join(" | ".join(cell.text for cell in row.cells) for row in table.rows)
                    except ImportError:
                        raise HTTPException(status_code=503, detail="DOCX extraction is unavailable. Install backend requirements and restart the server.")
                else:
                    raise HTTPException(status_code=400, detail="Legacy .doc files are not supported. Save the document as .docx or PDF and upload again.")
        else:
            text = content.decode("utf-8", errors="ignore")
        
        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail=(
                    "No embedded text was found in this PDF. It may be a scanned/image-only document; "
                    "upload a searchable PDF or DOCX file."
                ),
            )
        
        # Truncate to prevent context overflow (keep first 15000 chars)
        if len(text) > 15000:
            text = text[:15000] + "\n\n[Document truncated...]"
        
        document = {
            "filename": file.filename,
            "size": len(content),
            "text": text,
            "metadata": extract_case_metadata(text),
            "char_count": len(text),
        }
        document_store.setdefault(conversation_id, []).append(document)
        
        logger.info(f"Document uploaded: {file.filename} ({len(text)} chars) for conversation {conversation_id}")
        
        return {
            "conversation_id": conversation_id,
            "filename": file.filename,
            "size": len(content),
            "char_count": len(text),
            "document_count": len(document_store[conversation_id]),
            "text": text,
            "status": "uploaded",
            "message": f"Document '{file.filename}' uploaded successfully. You can now ask questions about it."
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/document/{conversation_id}")
async def remove_document(conversation_id: str):
    """Remove uploaded document from context"""
    if conversation_id in document_store:
        del document_store[conversation_id]
        return {"status": "removed", "conversation_id": conversation_id}
    return {"status": "not_found", "conversation_id": conversation_id}


@router.post("/classify")
async def classify_issue(request: ChatRequest):
    """Classify legal issue"""
    try:
        result = await google_client.classify_legal_issue(request.message)
        return result
    except Exception as e:
        logger.error(f"Error classifying issue: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/extract-entities")
async def extract_entities(request: ChatRequest):
    """Extract entities from case text"""
    try:
        result = await google_client.extract_case_entities(request.message)
        return result
    except Exception as e:
        logger.error(f"Error extracting entities: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/suggestions")
async def get_suggestions(query: str = ""):
    """Get suggested questions"""
    suggestions = [
        "Can I get bail under IPC 420?",
        "What is the procedure for filing an FIR?",
        "How long does a criminal case usually take?",
        "What evidence is needed for a fraud case?",
        "What are my rights as an accused?",
        "How do I apply for anticipatory bail?"
    ]
    return {"suggestions": suggestions}
